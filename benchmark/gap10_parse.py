#!/usr/bin/env python3
"""Gap 10b: parse the rendered vendor documents back into a corpus.

Three arms, dumbest first:

  naive       raw text dump (fitz page text / docx paragraphs — tables
              silently lost / xlsx tab-joined rows), then one generic regex
              segmenter over the text. The quick-script baseline.
  structured  format-native parsing: pdfplumber tables (tanager),
              heading-hierarchy walk of the docx body (kestrel), column-
              aware sheet reads (orrery), font-based line reconstruction of
              the two-column prose PDF (pellucid).
  docling     IBM Docling -> markdown -> one generic markdown segmenter.
              Runs in its own venv (its pin downgrades transformers):
              <venv-python> gap10_parse.py --arms docling

Recovered units are mapped back to question_ids by explicit Ref where the
format preserved one, else fuzzy match of the recovered question against
the 261-question catalog (scoring aid only — qrels need ids; a production
indexer wouldn't need the mapping). Unmappable units stay in the corpus as
distractors, exactly as they would in production. Output:
results/gap10/parsed-{arm}.jsonl + coverage/fidelity stats.

  python3 gap10_parse.py --arms naive,structured
"""
import argparse
import difflib
import json
import re
from collections import defaultdict
from pathlib import Path

HERE = Path(__file__).parent
DOCS = HERE / "rfq-docs"
OUT = HERE / "results" / "gap10"

REF_RE = re.compile(r"\b([A-Z]{2,4}-\d{2}\.\d+)\b")
MARKER_RE = re.compile(
    r"(?m)(\b[A-Z]{2,4}-\d{2}\.\d+\b|^Q\d+\.\s|\bQ\d+\.\s|^\d+\.\d+\.\d+\s)")

FILES = {
    "tanager_geo": "tanager_geo.pdf",
    "kestrel_cloud": "kestrel_cloud.docx",
    "orrery_suite": "orrery_suite.xlsx",
    "pellucid_data": "pellucid_data.pdf",
}


def norm(s):
    return " ".join(
        "".join(c.lower() if c.isalnum() or c.isspace() else " " for c in s).split())


# ------------------------------------------------------------------ naive arm
def naive_units(tenant, path):
    if path.suffix == ".pdf":
        import fitz
        text = "\n".join(p.get_text() for p in fitz.open(path))
    elif path.suffix == ".docx":
        from docx import Document
        text = "\n".join(p.text for p in Document(str(path)).paragraphs)
    else:
        from openpyxl import load_workbook
        wb = load_workbook(str(path), read_only=True)
        text = "\n".join(
            "\t".join(str(c) for c in row if c is not None)
            for ws in wb for row in ws.iter_rows(values_only=True))
    marks = list(MARKER_RE.finditer(text))
    units = []
    for i, m in enumerate(marks):
        seg = text[m.start():marks[i + 1].start() if i + 1 < len(marks) else len(text)]
        ref = REF_RE.match(seg.strip())
        qm = seg.find("?")
        if 0 < qm < 600:
            qtext, atext = seg[:qm + 1], seg[qm + 1:]
        else:
            qtext, atext = seg.split("\n", 1) if "\n" in seg else (seg, "")
        qtext = re.sub(r"^[A-Z]{2,4}-\d{2}\.\d+\s*|^Q\d+\.\s*|^\d+\.\d+\.\d+\s*",
                       "", qtext.strip())
        units.append((ref.group(1) if ref else None,
                      " ".join(qtext.split()), " ".join(atext.split())))
    return units


# ------------------------------------------------------------- structured arm
def structured_pdf_tables(path):
    import pdfplumber
    units = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                for row in table:
                    cells = [(c or "").strip() for c in row]
                    if len(cells) < 3 or cells[0] == "Ref":
                        continue
                    cells[0] = re.sub(r"\s+", "", cells[0])  # refs wrap in-cell
                    if REF_RE.fullmatch(cells[0]):
                        units.append((cells[0], " ".join(cells[1].split()),
                                      " ".join(cells[2].split())))
                    elif units:  # row continued across a page break
                        ref, q, a = units[-1]
                        units[-1] = (ref, (q + " " + cells[1]).strip(),
                                     (a + " " + cells[2]).strip())
    return units


def structured_docx(path):
    from docx.document import Document as DocxDoc
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx import Document
    doc = Document(str(path))
    units, control, cur = [], None, None

    def flush():
        if cur and cur[1]:
            units.append((cur[0], cur[1], " ".join(" ".join(cur[2]).split())))

    for el in doc.element.body:
        if el.tag.endswith("}p"):
            p = Paragraph(el, doc)
            style, text = p.style.name, p.text.strip()
            if not text:
                continue
            if style == "Heading 2":
                m = REF_RE.search(text + ".0")  # "3.2 ACC-02" has no sub part
                control = re.search(r"[A-Z]{2,4}-\d{2}", text)
                control = control.group(0) if control else None
            elif style == "Heading 3":
                flush()
                num = re.match(r"^(\d+)\.(\d+)\.(\d+)\s+(.*)", text)
                sub, q = (num.group(3), num.group(4)) if num else (None, text)
                ref = f"{control}.{sub}" if control and sub else None
                cur = (ref, " ".join(q.split()), [])
            elif style.startswith("Heading"):
                flush()
                cur = None
            elif cur:
                cur[2].append(text)
        elif el.tag.endswith("}tbl") and cur:
            for cell_text in re.findall(r"<w:t[^>]*>([^<]*)</w:t>", el.xml):
                cur[2].append(cell_text)
    flush()
    return units


def structured_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(str(path), read_only=True)
    units = []
    for ws in wb:
        if ws.title == "Instructions":
            continue
        started = False
        for row in ws.iter_rows(values_only=True):
            a, b, c = (row + (None,) * 3)[:3]
            if a == "Ref":
                started = True
                continue
            if not started:
                continue
            if a and REF_RE.fullmatch(str(a)):
                units.append((str(a), " ".join(str(b).split()), " ".join(str(c).split())))
            elif c and units:  # continuation row: response spilled over
                ref, q, ans = units[-1]
                units[-1] = (ref, q, ans + " " + " ".join(str(c).split()))
    return units


def structured_prose_pdf(path):
    import pdfplumber
    units, cur = [], None

    def flush():
        if cur:
            units.append((None, " ".join(cur[0].split()), " ".join(cur[1].split())))

    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            mid, h = page.width / 2, page.height
            lines = defaultdict(list)
            for ch in page.chars:
                if ch["top"] > h - 60:          # footer
                    continue
                col = 0 if ch["x0"] < mid else 1
                lines[(col, round(ch["top"] / 3))].append(ch)
            for key in sorted(lines):
                chs = sorted(lines[key], key=lambda c: c["x0"])
                text = "".join(c["text"] for c in chs).strip()
                if not text:
                    continue
                bold = sum("Bold" in c["fontname"] for c in chs) > len(chs) / 2
                size = max(c["size"] for c in chs)
                if size > 10:                    # domain heading
                    flush()
                    cur = None
                elif bold and re.match(r"^Q\d+\.", text):
                    flush()
                    cur = [re.sub(r"^Q\d+\.\s*", "", text), ""]
                elif bold and cur is not None and not cur[1]:
                    cur[0] += " " + text         # wrapped question line
                elif cur is not None:
                    cur[1] += " " + text
    flush()
    return units


def structured_units(tenant, path):
    if tenant == "tanager_geo":
        return structured_pdf_tables(path)
    if tenant == "kestrel_cloud":
        return structured_docx(path)
    if tenant == "orrery_suite":
        return structured_xlsx(path)
    return structured_prose_pdf(path)


# ---------------------------------------------------------------- docling arm
def docling_units(tenant, path):
    from docling.document_converter import DocumentConverter
    md = DocumentConverter().convert(str(path)).document.export_to_markdown()
    units, cur = [], None

    def flush():
        if cur and cur[1]:
            units.append((cur[0], " ".join(cur[1].split()),
                          " ".join(" ".join(cur[2]).split())))

    for line in md.splitlines():
        line = line.strip()
        if not line or set(line) <= {"|", "-", " ", ":"}:
            continue
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if len(cells) >= 3 and cells[0] != "Ref":
                cells[0] = re.sub(r"\s+", "", cells[0])  # refs wrap in-cell
                if REF_RE.fullmatch(cells[0]):
                    flush()
                    cur = None
                    units.append((cells[0], " ".join(cells[1].split()),
                                  " ".join(cells[2].split())))
                elif units and not cells[0]:
                    ref, q, a = units[-1]
                    units[-1] = (ref, (q + " " + cells[1]).strip(),
                                 (a + " " + cells[2]).strip())
            elif cur:
                cur[2].append(" ".join(c for c in cells if c))
            continue
        heading = re.match(r"^#+\s+(.*)", line)
        body = heading.group(1) if heading else line
        body_clean = body.strip("*").strip()
        qm = re.match(r"^(?:\d+\.\d+\.\d+|Q\d+\.)\s*(.*)", body_clean)
        if (heading or body.startswith("**Q")) and qm and "?" in body_clean:
            flush()
            cur = (None, qm.group(1), [])
        elif heading:
            flush()
            cur = None
        elif cur is not None:
            cur[2].append(line.strip("*"))
    flush()
    return units


# ------------------------------------------------------------ map and report
def map_units(units, tenant, catalog, truth):
    ncat = {qid: norm(t) for qid, t in catalog.items()}
    mapped, junk, dupes = {}, [], 0
    for ref, qtext, atext in units:
        qid, score = None, 0.0
        if ref and ref in catalog:
            qid, score = ref, 2.0
        elif qtext:
            nq = norm(qtext)
            sm = difflib.SequenceMatcher(None, nq, "")
            for cand, nc in ncat.items():
                sm.set_seq2(nc)
                if sm.real_quick_ratio() > max(score, 0.55) and \
                        sm.quick_ratio() > max(score, 0.55):
                    r = sm.ratio()
                    if r > score:
                        qid, score = cand, r
            if score < 0.55:
                qid = None
        if qid and f"{tenant}-{qid.lower()}" in truth:
            key = f"{tenant}-{qid.lower()}"
            if key in mapped and mapped[key][0] >= score:
                dupes += 1
                continue
            if key in mapped:
                dupes += 1
            mapped[key] = (score, qtext, atext)
        else:
            junk.append((qtext, atext))
    return mapped, junk, dupes


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--arms", default="naive,structured")
    args = ap.parse_args()
    OUT.mkdir(parents=True, exist_ok=True)

    docs = [json.loads(l) for l in open(HERE / "docs-rfq.jsonl")]
    truth = {d["id"]: d for d in docs}
    catalog = {}
    for d in docs:
        catalog[d["question_id"]] = d["title"]

    arms = {"naive": naive_units, "structured": structured_units,
            "docling": docling_units}
    for arm in args.arms.split(","):
        lines, stats = [], {}
        for tenant, fname in FILES.items():
            units = arms[arm](tenant, DOCS / fname)
            mapped, junk, dupes = map_units(units, tenant, catalog, truth)
            n_true = sum(1 for d in docs if d["tenant"] == tenant)
            sim = [difflib.SequenceMatcher(None, norm(v[2]),
                                           norm(truth[k]["text"])).ratio()
                   for k, v in mapped.items()]
            stats[tenant] = dict(
                units=len(units), recovered=len(mapped), of=n_true,
                coverage=round(len(mapped) / n_true, 3),
                answer_sim=round(sum(sim) / len(sim), 3) if sim else 0,
                junk=len(junk), dupes=dupes)
            for k, (score, qtext, atext) in mapped.items():
                lines.append(json.dumps(
                    {"id": k, "tenant": tenant, "title": qtext, "text": atext,
                     "map_conf": round(min(score, 1.0), 3)}))
            for j, (qtext, atext) in enumerate(junk):
                lines.append(json.dumps(
                    {"id": f"{tenant}-junk-{j}", "tenant": tenant,
                     "title": (qtext or atext)[:160], "text": atext}))
        (OUT / f"parsed-{arm}.jsonl").write_text("\n".join(lines) + "\n")
        print(f"\n== {arm} ==")
        for tenant, s in stats.items():
            print(f"  {tenant:<14} units {s['units']:>4}  recovered "
                  f"{s['recovered']}/{s['of']} ({s['coverage']:.1%})  "
                  f"answer-sim {s['answer_sim']:.3f}  junk {s['junk']}  "
                  f"dupes {s['dupes']}")
        (OUT / f"stats-{arm}.json").write_text(json.dumps(stats, indent=1))


if __name__ == "__main__":
    main()
