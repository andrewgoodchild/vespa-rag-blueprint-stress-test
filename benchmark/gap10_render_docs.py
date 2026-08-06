#!/usr/bin/env python3
"""Gap 10a: render the questionnaire corpus into realistic messy documents.

The study's corpus is clean JSONL; real RFP responses arrive as PDFs, Word
files and spreadsheets. This renders the same 940 ground-truth answers into
four per-vendor document sets, each with a different real-world flavour of
mess, so the parse -> segment -> retrieve pipeline can be measured against
qrels we already hold. Answer and question text is verbatim — the mess is
structural only (formats, headings, tables, numbering, boilerplate).

  tanager_geo    PDF, formal tabular: per-domain sections, bordered
                 Ref|Question|Response tables that split across pages,
                 cover page, headers/footers.
  kestrel_cloud  DOCX, narrative: H1 domain / H2 control / H3 question
                 heading hierarchy with decimal numbering, answers as
                 paragraphs — except every 7th answer inside a 1x1 table
                 ("callout box"), the classic Word trap.
  orrery_suite   XLSX, SIG-style workbook: Instructions sheet, one sheet
                 per domain, merged title cells, Ref|Question|Response|
                 Comments columns, every 9th answer split across two rows
                 (continuation row with empty Ref).
  pellucid_data  PDF, two-column prose: no tables, no control IDs anywhere
                 — questions are bold run-in "Q17." paragraphs with
                 continuous numbering; the segmentation + mapping stress
                 case.

Every vendor also gets generic boilerplate (cover/instructions) — the
cosine-attractor material Exp 3 warned about. Output -> rfq-docs/.

  python3 gap10_render_docs.py
"""
import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (BaseDocTemplate, Frame, LongTable, PageBreak,
                                PageTemplate, Paragraph, Spacer, TableStyle)

HERE = Path(__file__).parent
OUT = HERE / "rfq-docs"
DATE = "August 6, 2026"

VENDOR_NAMES = {
    "tanager_geo": "Tanager Geospatial, Inc.",
    "kestrel_cloud": "Kestrel Cloud Systems",
    "orrery_suite": "Orrery Software",
    "pellucid_data": "Pellucid Data Ltd.",
}

BOILER = (
    "This document contains {vendor}'s responses to the security and "
    "compliance questionnaire. The information herein describes our security "
    "program, policies and technical controls, and is provided in confidence "
    "for vendor-assessment purposes only. Responses reflect the state of our "
    "security posture as of {date}. Our organization maintains a "
    "comprehensive information security program with controls covering "
    "access management, encryption, incident response, business continuity "
    "and vendor risk. Questions regarding these responses should be directed "
    "to our security and compliance team."
)


def by_domain(docs):
    dom = defaultdict(list)
    for d in docs:
        dom[d["control_id"].split("-")[0]].append(d)
    for v in dom.values():
        v.sort(key=lambda d: d["question_id"])
    return dict(sorted(dom.items()))


# ---------------------------------------------------------------- tanager PDF
def render_tanager(docs, path, vendor):
    styles = getSampleStyleSheet()
    cell = ParagraphStyle("cell", parent=styles["BodyText"], fontSize=8.5, leading=11)
    head = ParagraphStyle("head", parent=styles["Heading1"], fontSize=15)

    def decorate(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.drawString(0.8 * inch, 0.55 * inch, vendor)
        canvas.drawCentredString(letter[0] / 2, 0.55 * inch,
                                 "Confidential — Security Questionnaire Response")
        canvas.drawRightString(letter[0] - 0.8 * inch, 0.55 * inch,
                               f"Page {doc_.page}")
        canvas.restoreState()

    doc = BaseDocTemplate(str(path), pagesize=letter,
                          topMargin=0.9 * inch, bottomMargin=0.9 * inch)
    frame = Frame(0.8 * inch, 0.9 * inch, letter[0] - 1.6 * inch,
                  letter[1] - 1.8 * inch)
    doc.addPageTemplates([PageTemplate(id="p", frames=[frame], onPage=decorate)])

    story = [Spacer(1, 2.2 * inch),
             Paragraph(vendor, styles["Title"]),
             Paragraph("Security &amp; Compliance Questionnaire Response",
                       styles["Heading2"]),
             Paragraph(f"Prepared {DATE} — CONFIDENTIAL", styles["Normal"]),
             PageBreak(),
             Paragraph("About this document", head),
             Paragraph(BOILER.format(vendor=vendor, date=DATE), styles["BodyText"]),
             PageBreak()]

    for si, (dom, items) in enumerate(by_domain(docs).items(), 1):
        story.append(Paragraph(f"Section {si} — {dom} Controls", head))
        story.append(Spacer(1, 8))
        rows = [[Paragraph("<b>Ref</b>", cell), Paragraph("<b>Question</b>", cell),
                 Paragraph("<b>Response</b>", cell)]]
        for d in items:
            rows.append([Paragraph(d["question_id"], cell),
                         Paragraph(d["title"], cell),
                         Paragraph(d["text"], cell)])
        t = LongTable(rows, colWidths=[0.7 * inch, 2.0 * inch, 4.2 * inch],
                      repeatRows=1)
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.9, 0.9, 0.93)),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story += [t, PageBreak()]
    doc.build(story)


# ---------------------------------------------------------------- kestrel DOCX
def render_kestrel(docs, path, vendor):
    doc = Document()
    doc.add_heading(vendor, 0)
    doc.add_paragraph("Security & Compliance Questionnaire Response")
    doc.add_paragraph(f"Prepared {DATE} — Confidential")
    doc.add_heading("About this response", 1)
    doc.add_paragraph(BOILER.format(vendor=vendor, date=DATE))

    n_boxed = 0
    for si, (dom, items) in enumerate(by_domain(docs).items(), 1):
        doc.add_heading(f"{si}. {dom} Controls", 1)
        last_control, ci = None, 0
        for qi, d in enumerate(items):
            if d["control_id"] != last_control:
                ci += 1
                last_control = d["control_id"]
                doc.add_heading(f"{si}.{ci} {d['control_id']}", 2)
            sub = d["question_id"].split(".")[-1]
            doc.add_heading(f"{si}.{ci}.{sub} {d['title']}", 3)
            if qi % 7 == 3:  # every 7th answer lives inside a callout table
                cell = doc.add_table(rows=1, cols=1).cell(0, 0)
                cell.text = d["text"]
                n_boxed += 1
            else:
                doc.add_paragraph(d["text"])
    doc.save(str(path))
    return n_boxed


# ---------------------------------------------------------------- orrery XLSX
def render_orrery(docs, path, vendor):
    wb = Workbook()
    ws = wb.active
    ws.title = "Instructions"
    ws.merge_cells("A1:D1")
    ws["A1"] = f"{vendor} — Security Questionnaire Response"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A3"] = BOILER.format(vendor=vendor, date=DATE)
    ws["A3"].alignment = Alignment(wrap_text=True, vertical="top")
    ws.merge_cells("A3:D8")

    n_split = 0
    for dom, items in by_domain(docs).items():
        ws = wb.create_sheet(dom)
        ws.merge_cells("A1:D1")
        ws["A1"] = f"{dom} Controls — {vendor}"
        ws["A1"].font = Font(bold=True, size=12)
        for c, h in zip("ABCD", ("Ref", "Question", "Vendor Response", "Comments")):
            ws[f"{c}2"] = h
            ws[f"{c}2"].font = Font(bold=True)
        ws.column_dimensions["B"].width = 50
        ws.column_dimensions["C"].width = 80
        r = 3
        for qi, d in enumerate(items):
            ws.cell(r, 1, d["question_id"])
            ws.cell(r, 2, d["title"]).alignment = Alignment(wrap_text=True, vertical="top")
            if qi % 9 == 4 and ". " in d["text"][80:]:  # split answer across 2 rows
                cut = 80 + d["text"][80:].index(". ") + 1
                ws.cell(r, 3, d["text"][:cut]).alignment = Alignment(wrap_text=True, vertical="top")
                r += 1
                ws.cell(r, 3, d["text"][cut + 1:]).alignment = Alignment(wrap_text=True, vertical="top")
                n_split += 1
            else:
                ws.cell(r, 3, d["text"]).alignment = Alignment(wrap_text=True, vertical="top")
            if qi % 15 == 7:
                ws.cell(r, 4, "Refer to policy document.")
            r += 1
    wb.save(str(path))
    return n_split


# ---------------------------------------------------------------- pellucid PDF
def render_pellucid(docs, path, vendor):
    styles = getSampleStyleSheet()
    body = ParagraphStyle("pbody", parent=styles["BodyText"], fontSize=8.5,
                          leading=11.5, spaceAfter=5)
    q = ParagraphStyle("pq", parent=body, spaceBefore=4)
    head = ParagraphStyle("phead", parent=styles["Heading2"], fontSize=12,
                          spaceBefore=10)

    def decorate(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica-Oblique", 7.5)
        canvas.drawString(0.7 * inch, 0.5 * inch,
                          f"{vendor} — Security Posture Statement")
        canvas.drawRightString(letter[0] - 0.7 * inch, 0.5 * inch, str(doc_.page))
        canvas.restoreState()

    doc = BaseDocTemplate(str(path), pagesize=letter)
    w = (letter[0] - 1.4 * inch - 0.3 * inch) / 2
    frames = [Frame(0.7 * inch, 0.8 * inch, w, letter[1] - 1.6 * inch),
              Frame(0.7 * inch + w + 0.3 * inch, 0.8 * inch, w,
                    letter[1] - 1.6 * inch)]
    doc.addPageTemplates([PageTemplate(id="two", frames=frames, onPage=decorate)])

    story = [Paragraph(vendor, styles["Title"]),
             Paragraph("Security Posture Statement — Questionnaire Responses",
                       styles["Heading2"]),
             Paragraph(f"Prepared {DATE} — Confidential", body),
             Paragraph(BOILER.format(vendor=vendor, date=DATE), body)]
    qn = 0
    for dom, items in by_domain(docs).items():
        story.append(Paragraph(f"{dom} — Domain Overview", head))
        for d in items:
            qn += 1
            story.append(Paragraph(f"<b>Q{qn}. {d['title']}</b>", q))
            story.append(Paragraph(d["text"], body))
    doc.build(story)
    return qn


def main():
    docs = [json.loads(l) for l in open(HERE / "docs-rfq.jsonl")]
    OUT.mkdir(exist_ok=True)
    per = defaultdict(list)
    for d in docs:
        per[d["tenant"]].append(d)

    manifest = {}
    for tenant, renderer, ext in (
            ("tanager_geo", render_tanager, "pdf"),
            ("kestrel_cloud", render_kestrel, "docx"),
            ("orrery_suite", render_orrery, "xlsx"),
            ("pellucid_data", render_pellucid, "pdf")):
        path = OUT / f"{tenant}.{ext}"
        extra = renderer(per[tenant], path, VENDOR_NAMES[tenant])
        manifest[path.name] = {"tenant": tenant, "answers": len(per[tenant]),
                               "bytes": path.stat().st_size, "note": extra}
        print(f"{path.name:<22} {len(per[tenant])} answers, "
              f"{path.stat().st_size//1024} KB (note={extra})")
    (OUT / "MANIFEST.json").write_text(json.dumps(manifest, indent=1))


if __name__ == "__main__":
    main()
